from template import *
import subprocess
import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pandas as pd
import tkinter as tk
from tkinter.ttk import *
from tkinter import filedialog
from ttkthemes import ThemedTk

# function that generates the doc file from the information entered in the GUI
def generate():
    # get all the information from the entry text and generate the file
    lawyer = et_lawyer_text.get()
    client = et_client_text.get()
    service = et_service_text.get()
    compensation_value = et_compensation_text.get()
    deposit_value = et_deposit_text.get()
    refundable_deposit_value = et_refundable_text.get()
    nonrefundable_deposit_value = et_non_refundable_text.get()
    deposit_date = et_deposit_date_text.get()
    jurisdiction = et_jurisdiction_text.get()

    # start generating the document
    doc = docx.Document()
    # title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.font.bold = True

    # parties
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(parties % (lawyer, client))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # service
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(services % service)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # responsibilities
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(responsibilities)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # compensation
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation % compensation_value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation3)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation4)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # costs
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(costs)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # deposit
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(deposit % (deposit_value, deposit_date, refundable_deposit_value, nonrefundable_deposit_value))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # provisions
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(provisions % jurisdiction)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # effective date
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(effective_date)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # foregoing
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(foregoing)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # signatures
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(signatures)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # save the doc
    doc.save(client+" Contract.docx")

    # update the status bar
    lb_status.config(text="Generated file " + client + " Contract.docx")


# function that generates all the files from the information entered in the Excel file
def generate_excel():
    filename = filedialog.askopenfilename(filetypes=(("Excel Files", "*.xlsx"), ("All files", "*.*")))
    # open the excel document and extract the data
    df = pd.read_excel(filename)

    lawyer = df["Attorney"]
    client = df["Client"]
    service = df["Service"]
    compensation_value = df["Compensation Value"]
    deposit_value = df["Deposit Value"]
    refundable_deposit_value = df["Refundable Deposit Value"]
    nonrefundable_deposit_value = df["Nonrefundable Deposit Value"]
    deposit_date = df["Deposit Date"]
    jurisdiction = df["Jurisdiction"]

    # start generating the document in a for. one document for each entry in the excel
    for i in range(len(lawyer)):
        doc = docx.Document()
        # title
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(20)
        run.font.bold = True

        # parties
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(parties % (lawyer[i], client[i]))
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # service
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(services % service[i])
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # responsibilities
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(responsibilities)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # compensation
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(compensation % compensation_value[i])
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(compensation2)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(compensation3)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(compensation4)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # costs
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(costs)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # deposit
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(
            deposit % (deposit_value[i], deposit_date[i], refundable_deposit_value[i], nonrefundable_deposit_value[i]))
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # provisions
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(provisions % jurisdiction[i])
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # effective date
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(effective_date)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # foregoing
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(foregoing)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # signatures
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(signatures)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        # save the doc, with the client as file name
        doc.save(client[i] + " Contract.docx")

        lb_status.config(text="Generated %s files" % len(lawyer))


# function that will call word and open a contract that we have generated
def open_contract():
    filename = filedialog.askopenfilename(initialdir=os.getcwd(), filetypes=(("Word Files", "*.docx"), ("All files", "*.*")))
    # there is different behavior on Windows vs Linux
    if tk.sys.platform == "win32":
        os.startfile('%s' % filename)
    else:
        # we assume Mac since we do not support linux yet
        subprocess.call(["open", filename])


# function that will call excel to edit the excel file used to generate contracts
def modify_excel():
    filename = filedialog.askopenfilename(initialdir=os.getcwd(), filetypes=(("Excel Files", "*.xlsx"), ("All files", "*.*")))
    # there is different behavior on Windows vs Linux
    if tk.sys.platform == "win32":
        os.startfile('%s' % filename)
    else:
        # we assume Mac since we do not support linux yet
        subprocess.call(["open", filename])


# function that deletes a contract
def delete_contract():
    filename = filedialog.askopenfilename(filetypes=(("Word Files", "*.docx"), ("All files", "*.*")))
    os.remove('%s' % filename)


# main part
# create the window
window = ThemedTk(screenName="Contract Generator", theme="sriv")


window.title("Generate Docs")
# window.geometry("1400x700")
window.resizable(0, 0)

# with ttk we need to configure styles:
style = Style()
style.configure("TButton", font=("Arial", 12, 'bold'), width=25)
style.configure("TLabel", font=("Arial", 15), anchor=tk.W, width=30, foreground="darkblue")
style.configure("TEntry", font=("Arial", 15), anchor=tk.W)

title_image = tk.PhotoImage(file="doc-generator.gif")
lb = Label(master=window, image=title_image)
lb.grid(row=0, pady=20, columnspan=3)

# Add the lawyer label and textbox
lb_lawyer = Label(master=window, text="Lawyer Name")
lb_lawyer.grid(row=1, column=0, padx=10)

et_lawyer_text = tk.StringVar()
et_lawyer = Entry(master=window, width=30, textvariable=et_lawyer_text)
et_lawyer.grid(row=1, column=1)

# Add the client label and textbox
lb_client = Label(master=window, text="Client Name")
lb_client.grid(row=2, column=0)

et_client_text = tk.StringVar()
et_client = Entry(master=window, width=30, textvariable=et_client_text)
et_client.grid(row=2, column=1)

# Add the service label and textbox
lb_service = Label(master=window, text="Service")
lb_service.grid(row=3, column=0)

et_service_text = tk.StringVar()
et_service = Entry(master=window, width=30, textvariable=et_service_text)
et_service.grid(row=3, column=1)

# compensation_value
lb_compensation = Label(master=window, text="Compensation Value")
lb_compensation.grid(row=4, column=0)

et_compensation_text = tk.StringVar()
et_compensation = Entry(master=window, width=30, textvariable=et_compensation_text)
et_compensation.grid(row=4, column=1)

# deposit_value
lb_deposit = Label(master=window, text="Deposit Value")
lb_deposit.grid(row=5, column=0)

et_deposit_text = tk.StringVar()
et_deposit = Entry(master=window, width=30, textvariable=et_deposit_text)
et_deposit.grid(row=5, column=1)

# refundable_deposit_value
lb_refundable = Label(master=window, text="Refundable Deposit Value")
lb_refundable.grid(row=6, column=0)

et_refundable_text = tk.StringVar()
et_refundable = Entry(master=window, width=30, textvariable=et_refundable_text)
et_refundable.grid(row=6, column=1)

# nonrefundable_deposit_value
lb_not_refundable = Label(master=window, text="Nonrefundable Deposit Value")
lb_not_refundable.grid(row=7, column=0)

et_non_refundable_text = tk.StringVar()
et_non_refundable = Entry(master=window, width=30, textvariable=et_non_refundable_text)
et_non_refundable.grid(row=7, column=1)

# deposit_date
lb_deposit_date = Label(master=window, text="Deposit Date")
lb_deposit_date.grid(row=8, column=0)

et_deposit_date_text = tk.StringVar()
et_deposit_date = Entry(master=window, width=30, textvariable=et_deposit_date_text)
et_deposit_date.grid(row=8, column=1)

# jurisdiction
lb_jurisdiction = Label(master=window, text="Jurisdiction")
lb_jurisdiction.grid(row=9, column=0)

et_jurisdiction_text = tk.StringVar()
et_jurisdiction = Entry(master=window, width=30, textvariable=et_jurisdiction_text)
et_jurisdiction.grid(row=9, column=1)

# status bar
lb_status = Label(master=window, text="", width=40)
lb_status.grid(row=10, column=0, columnspan=3, pady=20)

# Add the buttons
bt_generate = Button(master=window, text="Generate Contract", command=generate)
bt_generate.grid(row=1, column=2, padx=30)

bt_load_excel = Button(master=window, text="Generate from Excel", command=generate_excel)
bt_load_excel.grid(row=2, column=2, padx=30)

bt_load_excel = Button(master=window, text="Modify the Excel", command=modify_excel)
bt_load_excel.grid(row=3, column=2, padx=30)

bt_open_contract = Button(master=window, text="Open Contract", command=open_contract)
bt_open_contract.grid(row=4, column=2, padx=30)

bt_open_contract = Button(master=window, text="Delete Contract", command=delete_contract)
bt_open_contract.grid(row=5, column=2, padx=30)

bt_exit = Button(master=window, text="Exit", command=tk.sys.exit)
bt_exit.grid(row=11, column=2, pady=20)

# main loop
window.mainloop()